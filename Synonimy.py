from bs4 import BeautifulSoup
import requests
import docx
from collections import Counter
from docx.shared import RGBColor
from prettytable import PrettyTable
import re


class Synonyms:
    my_word = []

    def __init__(self, path):
        self.path = path

    @staticmethod
    def search_for_synonym(word):
        url = f"https://www.synonimy.pl/synonim/{word}"
        try:
            result = requests.get(url)
            if "Nie znaleziono wpisanego słowa " in result.text:
                return [word]
            doc = BeautifulSoup(result.text, "html.parser")
            words = doc.select("a.load_word")
            list_of_synonyms = [word.text for word in words]
            return list_of_synonyms
        except Exception as err:
            print("Coś poszło nie tak", err)

    def open_file(self):
        doc = docx.Document(self.path)
        text = []
        for _ in doc.paragraphs:
            text.append(_.text)
        return ' '.join(text).replace("(", "").replace(")", "")

    def change_my_choice(self, my: list):
        file = re.split(r'([,()*?!. ])', Synonyms.open_file(self))
        result = []
        p = 0
        z = [x.lower() for x in file]
        for _ in file:
            for i in range(len(my)):
                if _.lower() == my[i]:
                    if _.isalpha():
                        word = Synonyms.search_for_synonym(_)
                        if len(word) <= 1:
                            result.append("*" + _)
                        else:
                            if p == len(word):
                                p = 0
                            try:
                                if _[0].isupper():
                                    result.append("-" + word[p].capitalize())
                                else:
                                    result.append("-" + word[p])
                                    p += 1
                            except IndexError:
                                p = 0
                    else:
                        result.append(_)
                else:
                    result.append(_)

        return ''.join(result)

    def change_all(self, my: list = None):
        if my is None:
            my = []
        self.my_word = my
        file = re.split(r'([,()*?!. ])', Synonyms.open_file(self))
        list_of_words = []
        result = []

        for _ in file:
            if _.lower() not in list_of_words:
                list_of_words.append(_.lower())

        p = 0
        z = [x.lower() for x in file]
        for _ in file:
            for i in range(len(list_of_words)):
                if _.lower() == list_of_words[i]:
                    if z.count(_.lower()) >= 2:
                        if _.lower() in my:
                            result.append(_)
                        else:
                            if _.isalpha():
                                word = Synonyms.search_for_synonym(_)
                                if len(word) <= 1:
                                    result.append("*" + _)
                                else:
                                    if p == len(word):
                                        p = 0
                                    try:
                                        if _[0].isupper():
                                            result.append("-" + word[p].capitalize())
                                        else:
                                            result.append("-" + word[p])
                                        p += 1
                                    except IndexError:
                                        p = 0
                            else:
                                result.append(_)
                    else:
                        result.append(_)
        return ''.join(result)

    def show_statistics(self):
        red = '\033[31m'
        green = "\033[32m"
        yellow = "\033[33m"
        reset = "\033[0m"
        text = Synonyms.open_file(self).replace(",", "").replace(".", "").replace("(", "").replace(")", "") \
            .replace("?", "").replace("!", "").lower().split()
        dictio = {}
        for x, y in sorted(Counter(text).items(), key=lambda item: item[1], reverse=True): dictio[x] = y
        max_value = 0
        for x, y in dictio.items():
            z = round((y / len(text) * 100), 2)
            dictio[x] = (y, str(z) + '%')
            if y > max_value:
                max_value = y
        my_table = PrettyTable(["Słowo", "Ilość wystąpień", "Procent wystąpień"])
        for x, y in dictio.items():
            if y[0] < max_value // 2 or y[0] == 1:
                my_table.add_row([green + x + reset, green + str(y[0]) + reset, green + y[1] + reset])
            elif max_value // 2 <= y[0] < max_value:
                my_table.add_row([yellow + x + reset, yellow + str(y[0]) + reset, yellow + y[1] + reset])
            else:
                my_table.add_row([red + x + reset, red + str(y[0]) + reset, red + y[1] + reset])
        my_table.add_autoindex("Nr")
        print(my_table)

    def save_as_docx(self, o, file_name):
        doc = docx.Document()
        changed_file = o.split()
        paragraph = doc.add_paragraph("")

        for _ in changed_file:
            if _.startswith("*"):
                paragraph.add_run(_.replace("*", "") + " ").font.color.rgb = RGBColor(255, 0, 0)
            elif _.startswith("-"):
                paragraph.add_run(_.replace("-", "") + " ").bold = True
            else:
                paragraph.add_run(_ + " ")
        try:
            doc.save(f'{file_name}.docx')
            print("Zapisano")
        except:
            print("Wystąpił błąd")

/////////Menu/////////

path = input("Podaj ścieżkę do pliku docx: ")

synonyms = Synonyms(path)

while True:
    print("1. Zmiana wybranych słów.")
    print("2. Zmiana wszystkich słów.")
    print("3. Pokaż statystyki.")
    print("4. Wyjdź.")

    choice = input("Wybierz: ")

    if choice == "1":
        my_choice = input("Podaj słowo, które chcesz zmienić: ")
        result = synonyms.change_my_choice([my_choice])
        file_name = input("Podaj nazwę pliku docx do zapisania: ")
        synonyms.save_as_docx(result,file_name)

    elif choice == "2":
        my_list = input("Podaj słowa, które chcesz zachować (oddzielone przecinkami): ").split(",")
        file_name = input("Podaj nazwę pliku docx do zapisania: ")
        synonyms.save_as_docx(synonyms.change_all(my_list), file_name)

    elif choice == "3":
        synonyms.show_statistics()

    elif choice == "4":
        break

    else:
        print("Nieprawidłowa opcja. Spróbuj ponownie.")
